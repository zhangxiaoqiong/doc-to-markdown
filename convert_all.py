#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
综合文档转换程序（文件级串行处理）
加入智能前置检测：自动识别纯图片/扫描件，跳过无效文本提取，直接进入Vision视觉处理。
"""

import os
import sys
import subprocess
from pathlib import Path
import re
from datetime import datetime
import shutil
import tempfile

# Windows编码修复
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# ==================== 新增：扫描件智能探测 ====================
def is_scanned_pdf(file_path):
    """通过采样判断是否为纯图片/扫描件PDF"""
    if file_path.suffix.lower() != '.pdf':
        return False
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            text_length = 0
            # 采样前3页
            pages_to_check = min(3, len(pdf.pages))
            if pages_to_check == 0:
                return False
                
            for i in range(pages_to_check):
                text = pdf.pages[i].extract_text() or ""
                text_length += len(text.strip())
            
            # 如果平均每页提取的文本少于 50 个字符（通常只是页眉页脚或系统水印），判定为扫描件
            if text_length / pages_to_check < 50:
                return True
    except ImportError:
        print("    ⚠️ 提示: 缺少 pdfplumber 库，无法进行扫描件前置检测")
    except Exception as e:
        print(f"    ⚠️ 扫描件检测出现异常: {e}")
    return False
# ==========================================================

# 处理记录类
class ProcessRecord:
    """记录单个文件的处理信息"""
    def __init__(self, file_name):
        self.file_name = file_name
        self.source_size_mb = 0.0
        self.output_size_mb = 0.0
        self.methods = []  # 使用的处理方法列表
        self.has_images = False  # 是否检测到图片
        self.is_docx_with_images = False  # 是否是含有图片的DOCX
        self.has_anomalies = False  # 是否有异常
        self.status = "待处理"  # 完成/失败/警告
        self.notes = ""  # 备注

    def to_dict(self):
        """转换为字典用于Excel"""
        return {
            "文件名": self.file_name,
            "源文件大小(MB)": f"{self.source_size_mb:.2f}",
            "处理方法": " + ".join(self.methods) if self.methods else "未处理",
            "输出文件大小(MB)": f"{self.output_size_mb:.2f}",
            "检测到图片": "✅" if self.has_images else "❌",
            "异常标记": "⚠️" if self.has_anomalies else "✅",
            "状态": self.status,
            "备注": self.notes
        }

def get_failed_log_path(output_dir):
    """获取失败日志文件路径"""
    return Path(output_dir) / "failed.log"

def is_in_failed_log(output_dir, file_name):
    """检查文件是否在failed.log中"""
    failed_log = get_failed_log_path(output_dir)
    if not failed_log.exists():
        return False

    try:
        with open(failed_log, 'r', encoding='utf-8') as f:
            content = f.read()
            return file_name in content
    except:
        return False

def is_file_complete(output_dir, source_file_path):
    file_stem = source_file_path.stem
    output_file = Path(output_dir) / (file_stem + '.md')

    if is_in_failed_log(output_dir, source_file_path.name):
        return False

    if output_file.exists():
        return True

    return False

def log_failed(output_dir, file_name, error_msg):
    """记录失败的文件到failed.log"""
    failed_log = get_failed_log_path(output_dir)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(failed_log, 'a', encoding='utf-8') as f:
        f.write(f"[{timestamp}] {file_name} | {error_msg}\n")

def remove_from_failed_log(output_dir, file_name):
    """从failed.log中移除文件（处理成功后）"""
    failed_log = get_failed_log_path(output_dir)
    if not failed_log.exists():
        return

    try:
        with open(failed_log, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        filtered_lines = [line for line in lines if file_name not in line]

        with open(failed_log, 'w', encoding='utf-8') as f:
            f.writelines(filtered_lines)
    except:
        pass

def run_convert_docs_single_file(input_dir, output_dir, file_path):
    """处理单个文件：提取内容 + Claude转换"""
    cmd = [
        sys.executable,
        "convert_docs.py",
        "--input", str(input_dir),
        "--output", str(output_dir),
        file_path.name
    ]

    result = subprocess.run(cmd, capture_output=True, text=True)

    if result.returncode == 0:
        return True, None
    else:
        error_msg = result.stderr if result.stderr else result.stdout
        return False, error_msg

def has_actual_images_in_docx(file_path):
    """使用python-docx API精准检测DOCX中的实际图片对象"""
    if file_path.suffix.lower() != '.docx':
        return False, 0, []

    try:
        from docx import Document
        doc = Document(file_path)
        image_count = 0
        image_locations = []

        for para_idx, para in enumerate(doc.paragraphs):
            for run in para.runs:
                if 'drawing' in run._element.xml.lower():
                    image_count += 1
                    image_locations.append(f"段落{para_idx+1}中的图片")

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        for run in para.runs:
                            if 'drawing' in run._element.xml.lower():
                                image_count += 1
                                image_locations.append(f"表格{table_idx+1}行{row_idx+1}列{cell_idx+1}的图片")

        return image_count > 0, image_count, image_locations
    except Exception as e:
        return False, 0, []

def check_needs_vision_single_file(input_dir, output_dir, file_path):
    """检查单个文件是否需要进一步处理"""
    file_stem = file_path.stem
    md_file = Path(output_dir) / (file_stem + '.md')

    if not md_file.exists():
        return 'ok', None

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    if file_path.suffix.lower() == '.pdf':
        source_size_mb = file_path.stat().st_size / (1024 * 1024)
        if source_size_mb > 15:
            return 'split', f'PDF文件大于15MB，需要分割处理（当前{source_size_mb:.1f}MB）'

    source_size_mb = file_path.stat().st_size / (1024 * 1024)
    output_size_mb = md_file.stat().st_size / (1024 * 1024)
    size_ratio = output_size_mb / source_size_mb if source_size_mb > 0 else 0

    if size_ratio < 0.1 and source_size_mb > 0.5:
        return 'vision', f'输出过小（仅{size_ratio*100:.1f}%），可能提取失败'

    if '此部分在第' in content or '[此部分内容在' in content:
        return 'vision', '检测到分割占位符，输出可能不完整'

    text_length = len(content.strip())
    if text_length < 500:
        return 'vision', f'输出太短（仅{text_length}字符），可能提取失败'

    if file_path.suffix.lower() == '.docx':
        has_images, image_count, image_locations = has_actual_images_in_docx(file_path)
        if has_images:
            locations_str = "; ".join(image_locations[:3])
            if len(image_locations) > 3:
                locations_str += f" 等（共{image_count}张）"
            return 'ok', f'DOCX含有{image_count}张嵌入图片（{locations_str}），需手工检查'

    return 'ok', None

def clean_markdown_pollution(md_file_path):
    """清理多来源的markdown污染标记"""
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        original = content

        if content.startswith('```markdown\n'):
            content = content[len('```markdown\n'):]
        elif content.startswith('```markdown'):
            content = content[len('```markdown'):]
            if content.startswith('\n'):
                content = content[1:]

        if content.endswith('\n```'):
            content = content[:-4]
        elif content.endswith('\n```\n'):
            content = content[:-5]
        elif content.endswith('```'):
            content = content[:-3]

        lines = content.split('\n')
        cleaned_lines = []
        skip_next = False

        for i, line in enumerate(lines):
            if line.strip() == '```markdown':
                skip_next = False
                continue
            if line.strip() == '```' and i > 0 and i < len(lines) - 1:
                if cleaned_lines and cleaned_lines[-1].strip() != '':
                    continue
            cleaned_lines.append(line)

        if cleaned_lines:
            content = '\n'.join(cleaned_lines)

        content = content.strip()

        if content != original:
            with open(md_file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        return False
    except Exception as e:
        print(f"    ⚠️  清理污染标记失败: {e}")
        return False

def run_convert_pdf_vision_single_file(input_dir, output_dir, file_path):
    """对单个PDF或DOCX文件用Vision API重新处理"""
    if file_path.suffix.lower() == '.pdf':
        cmd = [
            sys.executable,
            "convert_split_pdf_v2.py",
            file_path.name,
            "--input", str(input_dir),
            "--output", str(output_dir)
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode == 0:
            return True, None
        else:
            error_msg = result.stderr if result.stderr else result.stdout
            return False, error_msg

    elif file_path.suffix.lower() == '.docx':
        return False, "DOCX文件含有流程/图片，无法自动处理。建议用Word打开手工检查。"

    return False, "未知文件类型"

def run_fix_markdown_single_file(output_dir, file_path):
    """校对单个Markdown文件的OCR错误"""
    file_stem = file_path.stem
    md_file = Path(output_dir) / (file_stem + '.md')

    if not md_file.exists():
        return True, None 

    temp_dir = None
    try:
        temp_dir = Path(tempfile.mkdtemp())
        temp_md_file = temp_dir / (file_stem + '.md')

        shutil.copy(md_file, temp_md_file)

        cmd = [
            sys.executable,
            "fix_markdown_with_claude.py",
            "--dir", str(temp_dir)
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode == 0:
            shutil.copy(temp_md_file, md_file)
            return True, None
        else:
            error_msg = result.stderr if result.stderr else result.stdout
            return False, error_msg

    except Exception as e:
        return False, f"校对失败: {str(e)}"

    finally:
        if temp_dir and temp_dir.exists():
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

def save_raw_pdf_output(output_dir, file_path):
    """保存PDF步骤1的原始输出为_step1_raw.md"""
    if file_path.suffix.lower() != '.pdf':
        return False

    try:
        file_stem = file_path.stem
        md_file = Path(output_dir) / (file_stem + '.md')
        raw_file = Path(output_dir) / (file_stem + '_step1_raw.md')

        if md_file.exists() and not raw_file.exists():
            shutil.copy(md_file, raw_file)
            return True
    except Exception as e:
        print(f"    ⚠️  保存原始输出失败: {e}")

    return False

def process_single_file(input_dir, output_dir, file_path):
    """处理单个文件的完整流程（4步分流版）"""
    file_name = file_path.name
    record = ProcessRecord(file_name)
    record.source_size_mb = file_path.stat().st_size / (1024 * 1024)

    print(f"\n==================================================")
    print(f"📄 开始处理: {file_name}")
    
    # ----------------------------------------------------
    # 新增逻辑：前置检测扫描件，智能分流
    # ----------------------------------------------------
    is_scanned = False
    if file_path.suffix.lower() == '.pdf':
        print(f"  [前置检测] 扫描文件内容特征... ", end="", flush=True)
        is_scanned = is_scanned_pdf(file_path)
        if is_scanned:
            print("判定为 [纯图片/扫描件]")
            record.notes = "系统识别为纯图片/扫描件"
        else:
            print("判定为 [原生可解析文档]")

    status = 'ok'
    reason = None

    if is_scanned:
        # 如果是扫描件，直接跳过 Step 1 和 Step 2，进入 Vision
        print(f"  ⚠️  [智能跳过] 已跳过文本提取(步骤1/2)，直接进入视觉识别(步骤3)")
        update_inventory_excel(output_dir, file_name, 'step1', '跳过', '纯扫描件')
        update_inventory_excel(output_dir, file_name, 'step2', '跳过', '无需质检')
        status = 'vision'
        reason = '纯图片/扫描件，需要视觉大模型直接处理'
        record.methods.append("skip_to_vision")
        
    else:
        # 正常流程：执行步骤 1
        print(f"  [步骤1] 文本提取与转换...")
        update_inventory_excel(output_dir, file_name, 'step1', '进行中')
        record.methods.append("convert_docs")
        success, error = run_convert_docs_single_file(input_dir, output_dir, file_path)
        
        if not success:
            record.status = "失败"
            record.notes = f"步骤1转换失败"
            update_inventory_excel(output_dir, file_name, 'step1', '失败', error[:100] if error else "未知错误")
            update_inventory_excel(output_dir, file_name, 'final_status', '失败', '步骤1转换失败')
            return False, f"步骤1转换失败: {error}", record

        if file_path.suffix.lower() == '.pdf':
            save_raw_pdf_output(output_dir, file_path)

        update_inventory_excel(output_dir, file_name, 'step1', '完成')

        # 正常流程：执行步骤 2
        print(f"  [步骤2] 质量检测...")
        update_inventory_excel(output_dir, file_name, 'step2', '进行中')
        status, reason = check_needs_vision_single_file(input_dir, output_dir, file_path)
        update_inventory_excel(output_dir, file_name, 'step2', status, reason or "")

    # ----------------------------------------------------
    # 处理分支状态（Vision重处理或跳过）
    # ----------------------------------------------------
    if status == 'split':
        record.methods.append("split")
        print(f"  [步骤3] 大文件分割处理 (原因: {reason})...")
        update_inventory_excel(output_dir, file_name, 'step3', '进行中', reason or "")
        success, error = run_convert_pdf_vision_single_file(input_dir, output_dir, file_path)
        if not success:
            record.status = "失败"
            record.notes = f"步骤3分割处理失败"
            update_inventory_excel(output_dir, file_name, 'step3', '失败', error[:100] if error else "未知错误")
            update_inventory_excel(output_dir, file_name, 'final_status', '失败', '步骤3分割处理失败')
            return False, f"步骤3分割处理失败: {error}", record

        update_inventory_excel(output_dir, file_name, 'step3', '完成')

    elif status == 'vision':
        if "skip_to_vision" not in record.methods:
            record.methods.append("vision")
            record.has_anomalies = True
            
        print(f"  [步骤3] Vision 视觉处理 (原因: {reason})...")
        update_inventory_excel(output_dir, file_name, 'step3', '进行中', reason or "")
        success, error = run_convert_pdf_vision_single_file(input_dir, output_dir, file_path)
        if not success:
            record.status = "失败"
            record.notes = f"步骤3 Vision处理失败"
            update_inventory_excel(output_dir, file_name, 'step3', '失败', error[:100] if error else "未知错误")
            update_inventory_excel(output_dir, file_name, 'final_status', '失败', '步骤3 Vision处理失败')
            return False, f"步骤3 Vision处理失败: {error}", record

        update_inventory_excel(output_dir, file_name, 'step3', '完成')

        file_stem = file_path.stem
        md_file = Path(output_dir) / (file_stem + '.md')
        if md_file.exists():
            clean_markdown_pollution(md_file)

    else:  # status == 'ok'
        update_inventory_excel(output_dir, file_name, 'step3', 'ok')
        print(f"  [步骤3] 质量良好，无需视觉重置。")

        if file_path.suffix.lower() == '.docx':
            has_images, image_count, image_locations = has_actual_images_in_docx(file_path)
            if has_images:
                record.is_docx_with_images = True
                record.has_anomalies = True
                record.has_images = True
                record.status = "警告"
                locations_str = "; ".join(image_locations[:2])
                if len(image_locations) > 2:
                    locations_str += f" 等"
                record.notes = f"DOCX含有{image_count}张图片（{locations_str}），需手工补充"
                print(f"  ⚠️  [警告] DOCX文件含有{image_count}张嵌入图片")
                
                file_stem = file_path.stem
                md_file = Path(output_dir) / (file_stem + '.md')
                if md_file.exists():
                    record.output_size_mb = md_file.stat().st_size / (1024 * 1024)

                update_inventory_excel(output_dir, file_name, 'step4', 'ok')
                update_inventory_excel(output_dir, file_name, 'final_status', '警告', f'DOCX含有{image_count}张图片')
                return True, None, record 

    # ----------------------------------------------------
    # 步骤 4：全文内容校对与润色
    # ----------------------------------------------------
    print(f"  [步骤4] 大模型 OCR 校对与润色...")
    update_inventory_excel(output_dir, file_name, 'step4', '进行中')
    record.methods.append("fix_markdown")
    success, error = run_fix_markdown_single_file(output_dir, file_path)
    
    if not success:
        record.status = "失败"
        record.notes = f"步骤4校对失败"
        update_inventory_excel(output_dir, file_name, 'step4', '失败', error[:100] if error else "未知错误")
        update_inventory_excel(output_dir, file_name, 'final_status', '失败', '步骤4校对失败')
        return False, f"步骤4校对失败: {error}", record

    file_stem = file_path.stem
    md_file = Path(output_dir) / (file_stem + '.md')
    if md_file.exists():
        clean_markdown_pollution(md_file)

    update_inventory_excel(output_dir, file_name, 'step4', '完成')

    # 记录最终状态
    if md_file.exists():
        record.output_size_mb = md_file.stat().st_size / (1024 * 1024)
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        if '![' in content:
            record.has_images = True

    record.status = "完成"
    update_inventory_excel(output_dir, file_name, 'final_status', '完成', '处理成功')
    return True, None, record

def get_inventory_excel_path(output_dir):
    """获取清单Excel的路径"""
    return Path(output_dir) / "inventory.xlsx"

def load_openpyxl():
    """导入openpyxl库"""
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        return Workbook, load_workbook, Font, PatternFill, Alignment, Border, Side
    except ImportError:
        print("提示：需要安装 openpyxl 库才能生成 Excel")
        return None, None, None, None, None, None, None

def generate_inventory_excel(output_dir, source_files):
    """生成初始清单Excel"""
    Workbook, load_workbook, Font, PatternFill, Alignment, Border, Side = load_openpyxl()
    if Workbook is None:
        return False

    try:
        inventory_path = get_inventory_excel_path(output_dir)
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "处理进度"

        headers = [
            "文件名", "文件大小(MB)", "源文件路径",
            "步骤1(转换)", "步骤2(检测)", "步骤3(处理)", "步骤4(校对)",
            "最终状态", "最后更新时间", "备注"
        ]
        worksheet.append(headers)

        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

        for file_path in source_files:
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            row_data = [
                file_path.name, f"{file_size_mb:.2f}", str(file_path),
                "待处理", "-", "-", "-", "待处理",
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"), ""
            ]
            worksheet.append(row_data)

            row_num = worksheet.max_row
            status_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")

            for col_num, cell in enumerate(worksheet[row_num], 1):
                cell.fill = status_fill
                cell.border = thin_border
                if col_num == 2:
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                else:
                    cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

        column_widths = [30, 12, 35, 15, 15, 15, 15, 12, 20, 35]
        for i, width in enumerate(column_widths, 1):
            worksheet.column_dimensions[chr(64 + i)].width = width

        worksheet.freeze_panes = "A2"
        workbook.save(str(inventory_path))
        return True
    except Exception as e:
        print(f"⚠️  生成初始清单失败: {e}")
        return False

def update_inventory_excel(output_dir, file_name, step, status, reason=""):
    """实时更新清单Excel中的单个文件的进度"""
    Workbook, load_workbook, Font, PatternFill, Alignment, Border, Side = load_openpyxl()
    if load_workbook is None:
        return False

    try:
        inventory_path = get_inventory_excel_path(output_dir)
        if not inventory_path.exists():
            return False

        workbook = load_workbook(str(inventory_path))
        worksheet = workbook.active

        row_num = None
        for row in range(2, worksheet.max_row + 1):
            if worksheet.cell(row, 1).value == file_name:
                row_num = row
                break

        if row_num is None:
            return False

        step_col_map = {'step1': 4, 'step2': 5, 'step3': 6, 'step4': 7, 'final_status': 8}

        if step in step_col_map:
            col_num = step_col_map[step]
            worksheet.cell(row_num, col_num).value = status

        worksheet.cell(row_num, 9).value = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        if reason:
            current_notes = worksheet.cell(row_num, 10).value or ""
            if current_notes and reason not in current_notes:
                worksheet.cell(row_num, 10).value = f"{current_notes} | {reason}"
            else:
                worksheet.cell(row_num, 10).value = reason

        if step == 'final_status':
            status_fill_map = {
                "完成": PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"),
                "警告": PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"),
                "失败": PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"),
                "待处理": PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid"),
            }
            fill = status_fill_map.get(status, PatternFill())

            for col in range(1, 11):
                worksheet.cell(row_num, col).fill = fill

        workbook.save(str(inventory_path))
        return True
    except Exception as e:
        return False

def get_file_status_from_inventory(output_dir, file_name):
    """从inventory.xlsx中读取文件的完成状态和当前步骤"""
    Workbook, load_workbook, Font, PatternFill, Alignment, Border, Side = load_openpyxl()
    if load_workbook is None:
        return False, 'unknown', None

    try:
        inventory_path = get_inventory_excel_path(output_dir)
        if not inventory_path.exists():
            return False, 'step1', None

        workbook = load_workbook(str(inventory_path))
        worksheet = workbook.active

        row_num = None
        for row in range(2, worksheet.max_row + 1):
            if worksheet.cell(row, 1).value == file_name:
                row_num = row
                break

        if row_num is None:
            return False, 'step1', None

        final_status = worksheet.cell(row_num, 8).value

        if final_status == "完成":
            return True, 'all', final_status
        elif final_status == "失败":
            return False, 'unknown', final_status
        elif final_status == "警告":
            return True, 'all', final_status
        else:
            step1 = worksheet.cell(row_num, 4).value
            step2 = worksheet.cell(row_num, 5).value
            step3 = worksheet.cell(row_num, 6).value
            step4 = worksheet.cell(row_num, 7).value

            if step1 in ["完成", "跳过"]:
                if step2 in ["完成", "ok", "split", "vision", "跳过"]:
                    if step3 == "完成":
                        if step4 == "完成":
                            return True, 'all', final_status
                        else:
                            return False, 'step4', final_status
                    else:
                        return False, 'step3', final_status
                else:
                    return False, 'step2', final_status
            else:
                return False, 'step1', final_status
    except Exception as e:
        return False, 'unknown', None

def generate_record_excel(output_dir, records):
    """生成Excel处理记录表"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        return

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "处理记录"

    headers = ["文件名", "源文件大小(MB)", "处理方法", "输出文件大小(MB)", "检测到图片", "异常标记", "状态", "备注"]
    worksheet.append(headers)

    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    for cell in worksheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border

    status_fill_map = {
        "完成": PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"),
        "警告": PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"),
        "失败": PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"),
        "已跳过": PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid"),
    }

    for record in records:
        record_dict = record.to_dict()
        row_data = [
            record_dict["文件名"], record_dict["源文件大小(MB)"], record_dict["处理方法"],
            record_dict["输出文件大小(MB)"], record_dict["检测到图片"], record_dict["异常标记"],
            record_dict["状态"], record_dict["备注"]
        ]
        worksheet.append(row_data)

        row_num = worksheet.max_row
        status = record.status
        fill = status_fill_map.get(status, PatternFill())

        for col_num, cell in enumerate(worksheet[row_num], 1):
            cell.fill = fill
            cell.border = thin_border
            if col_num in [2, 4]:
                cell.alignment = Alignment(horizontal='right', vertical='center')
            else:
                cell.alignment = Alignment(horizontal='left', vertical='center')

    column_widths = [35, 15, 25, 15, 12, 10, 10, 30]
    for i, width in enumerate(column_widths, 1):
        worksheet.column_dimensions[chr(64 + i)].width = width

    excel_path = Path(output_dir) / "处理记录.xlsx"
    workbook.save(str(excel_path))

def main():
    if not os.getenv("ANTHROPIC_BASE_URL") or not os.getenv("ANTHROPIC_AUTH_TOKEN"):
        print("错误: 缺少必要的环境变量")
        sys.exit(1)

    import argparse
    parser = argparse.ArgumentParser(description='综合文档转换程序：智能分流、四步精细化处理')
    parser.add_argument('--input', '-i', default='知识库base', help='输入目录')
    parser.add_argument('--output', '-o', default='知识库md_v1.0', help='输出目录')
    args = parser.parse_args()

    input_dir = Path(args.input)
    output_dir = Path(args.output)

    if not input_dir.exists():
        print(f"错误：输入目录不存在 {input_dir}")
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)

    source_files = sorted([
        f for f in input_dir.glob('*')
        if f.suffix.lower() in ['.pdf', '.docx']
    ])

    if not source_files:
        print(f"错误：没有找到PDF或DOCX文件在 {input_dir}")
        sys.exit(1)

    print("="*60)
    print(f"📋 生成/加载处理清单...")
    print("="*60)

    inventory_path = get_inventory_excel_path(output_dir)
    if not inventory_path.exists():
        if generate_inventory_excel(output_dir, source_files):
            print(f"✓ 初始清单已生成: {inventory_path}")
    else:
        print(f"✓ 使用已存在的进度清单: {inventory_path}")

    print("\n" + "="*60)
    print(f"🔄 开始智能处理 {len(source_files)} 个文件")
    print("="*60)

    completed = 0
    skipped = 0
    failed = 0
    records = []

    for file_path in source_files:
        file_name = file_path.name
        is_complete, current_step, final_status = get_file_status_from_inventory(output_dir, file_name)

        if is_complete:
            print(f"\n[跳过] {file_name} - 已完成 ({final_status})")
            skipped += 1
            record = ProcessRecord(file_name)
            record.source_size_mb = file_path.stat().st_size / (1024 * 1024)
            file_stem = file_path.stem
            md_file = Path(output_dir) / (file_stem + '.md')
            if md_file.exists():
                record.output_size_mb = md_file.stat().st_size / (1024 * 1024)
            record.status = "已跳过"
            record.notes = "文件已处理过"
            records.append(record)
            continue

        if current_step == 'unknown':
            pass
        else:
            print(f"\n[断点续传] {file_name} - 准备从 {current_step} 继续...")

        success, error_msg, record = process_single_file(input_dir, output_dir, file_path)
        records.append(record)

        if success:
            print(f"✓ [处理成功] {file_name}")
            remove_from_failed_log(output_dir, file_name)
            completed += 1
        else:
            print(f"✗ [处理失败] {file_name}")
            if error_msg:
                print(f"   错误详情: {error_msg[:200]}")
            log_failed(output_dir, file_name, error_msg or "未知错误")
            failed += 1

    print("\n" + "="*60)
    print("🎉 全部任务运行完毕！")
    print(f"  本次完成: {completed}")
    print(f"  无需处理: {skipped}")
    print(f"  发生失败: {failed}")

    try:
        generate_record_excel(output_dir, records)
        print(f"\n✓ 最终处理记录表已生成: {output_dir}/处理记录.xlsx")
    except Exception as e:
        print(f"\n⚠️  Excel生成失败（不影响转换结果）: {e}")

    if failed > 0:
        print(f"\n✗ 提示：有 {failed} 个文件处理失败，详情请看 {get_failed_log_path(output_dir)}")
        sys.exit(1)
    print("="*60)

if __name__ == "__main__":
    main()
