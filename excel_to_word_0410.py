#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Excel to Word 转换工具 - 用document类型处理图片
"""

import openpyxl
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import logging
import base64
import time
import os
from anthropic import Anthropic
from PIL import Image

logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)


def compress_image(image_data: bytes, max_size_mb: float = 4.5, max_dimension: int = 8000) -> bytes:
    """压缩图片以满足API大小限制"""
    try:
        # 打开图片
        img = Image.open(io.BytesIO(image_data))

        # 转换为RGB（JPEG不支持透明度）
        if img.mode in ('RGBA', 'LA', 'P'):
            # 创建白色背景
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if 'A' in img.mode else None)
            img = background
        else:
            img = img.convert('RGB')

        # 先检查尺寸是否超限
        width, height = img.size
        if width > max_dimension or height > max_dimension:
            # 需要缩小尺寸
            scale = min(max_dimension / width, max_dimension / height, 1.0)
            new_width = int(width * scale)
            new_height = int(height * scale)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            logger.info(f"      图片尺寸超过限制，已缩放: {width}x{height} → {new_width}x{new_height}")

        # 循环压缩直到满足大小要求
        quality = 90
        max_size_bytes = int(max_size_mb * 1024 * 1024)

        while quality > 10:
            output = io.BytesIO()
            img.save(output, format='JPEG', quality=quality, optimize=True)
            compressed_data = output.getvalue()

            if len(compressed_data) <= max_size_bytes:
                logger.info(f"      图片压缩成功 - 原始: {len(image_data)/1024/1024:.2f}MB, 压缩后: {len(compressed_data)/1024:.1f}KB (质量: {quality}%)")
                return compressed_data

            quality -= 10

        # 如果质量调到最低还是太大，继续缩小尺寸
        width, height = img.size
        while width > 1000 or height > 1000:
            scale = min(1000 / width, 1000 / height, 1.0)
            width = int(width * scale)
            height = int(height * scale)
            resized = img.resize((width, height), Image.Resampling.LANCZOS)
            output = io.BytesIO()
            resized.save(output, format='JPEG', quality=80, optimize=True)
            compressed_data = output.getvalue()
            if len(compressed_data) <= max_size_bytes:
                logger.info(f"      图片二次缩放成功 - 原始: {len(image_data)/1024/1024:.2f}MB, 压缩后: {len(compressed_data)/1024:.1f}KB (尺寸: {width}x{height})")
                return compressed_data

        logger.warning(f"      图片无法压缩到要求大小，返回原始数据")
        return image_data

    except Exception as e:
        logger.warning(f"      图片压缩异常: {str(e)}, 返回原始数据")
        return image_data



def analyze_image_with_question(image_data: bytes, question: str, row_idx: int, existing_answer: str = "", max_retries: int = 3) -> tuple:
    """使用Vision API分析图片，返回(完整内容, 图片标题)元组

    完整内容：表格、流程图等所有文字内容的详细解析
    图片标题：一句话说明这是什么
    """
    import random

    client = Anthropic(
        api_key=os.getenv("ANTHROPIC_AUTH_TOKEN"),
        base_url=os.getenv("ANTHROPIC_BASE_URL")
    )

    # 统一的prompt：无论答案是否存在，都要求完整内容+一句话标题
    prompt = f"""用户的问题是：{question}

请仔细查看这张图片，完成以下任务：

**任务1：完整还原图片中的所有文字内容**
- 如果是表格：按表格格式还原，包括所有行、列、标题、数据
- 如果是流程图：按流程步骤还原，包括每个节点、流向、条件判断
- 如果是分类标准：按分类层级还原，包括每个分类及其说明
- 保留原有的格式、编号、层级关系
- 使用markdown格式输出
- **不要**生成图片中不存在的内容
- **不要**提及任何错误、失败、无法读取的信息

**任务2：一句话标题**
- 用简洁的一句话说明这是什么（例如："出差原因分类标准表"、"员工报销流程图"、"酒店住宿费用标准表"）

请按以下格式返回，用"===SEPARATOR==="分隔两部分：
- 如果图片内容丰富（表格、流程图、多步骤指引等），详细内容用```markdown代码块包裹
- 如果图片只是一句话说明或简单标签，直接输出文字，不用markdown标记

示例（内容丰富）：
```markdown
# 出差申请单填写指引
## 一、申请指引 -APP端
路径：丰声—微服务...
```

===SEPARATOR===

出差申请单APP及PC端填写操作指引

示例（内容简单）：
深圳投单箱二维码

===SEPARATOR===

深圳投单箱二维码"""

    for attempt in range(max_retries):
        try:
            # 先压缩图片到可接受的大小
            image_data = compress_image(image_data)

            image_base64 = base64.standard_b64encode(image_data).decode("utf-8")

            # 使用 image 类型（企业代理的document类型对某些图片会返回"未检测到图片"）
            content = [
                {
                    "type": "text",
                    "text": prompt
                },
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": "image/jpeg",
                        "data": image_base64
                    }
                }
            ]

            message = client.messages.create(
                model="claude-opus-4-6",
                max_tokens=20000,
                messages=[{
                    "role": "user",
                    "content": content
                }]
            )

            # 提取回复文本
            response_text = ""
            for block in message.content:
                if hasattr(block, 'text'):
                    response_text = block.text
                    break

            if not response_text:
                return "[图片分析返回内容为空]", "[无法生成说明]"

            # 解析返回内容：分离完整内容和图片标题
            if '===SEPARATOR===' in response_text:
                parts = response_text.split('===SEPARATOR===', 1)
                full_content = parts[0].strip()
                title = parts[1].strip() if len(parts) > 1 else "图片"
                # 标题取第一行（去除前后空白）
                title_lines = [l.strip() for l in title.split('\n') if l.strip()]
                title = title_lines[0] if title_lines else "图片"
            else:
                # 降级方案：整个回复作为完整内容，生成一个简单标题
                full_content = response_text.strip()
                # 尝试提取第一行作为标题
                first_line = full_content.split('\n')[0]
                title = first_line[:50] if first_line else "图片"

            return full_content, title

        except Exception as e:
            error_msg = str(e)
            if attempt < max_retries - 1:
                wait_time = (2 ** (attempt + 1)) + random.uniform(0, 1)
                logger.warning(f"   行{row_idx}: 重试中 ({attempt+1}/{max_retries})...错误: {error_msg[:100]}")
                time.sleep(wait_time)
            else:
                logger.warning(f"   行{row_idx}: 图片分析失败 - {type(e).__name__}: {error_msg[:100]}")
                return "[分析失败]", "[分析失败]"

    return "[分析失败]", "[分析失败]"


def convert_excel_to_word(excel_path, output_path):
    """转换Excel到Word，使用Vision分析缺失答案中的图片"""

    excel_path = Path(excel_path)
    output_path = Path(output_path)

    if not excel_path.exists():
        logger.error(f"❌ 文件不存在: {excel_path}")
        return False

    try:
        logger.info(f"📖 处理: {excel_path.name}")

        # 加载Excel
        wb = openpyxl.load_workbook(excel_path)
        ws = wb[wb.sheetnames[0]]

        logger.info(f"   工作表: {ws.title}, 总行数: {ws.max_row}")

        # 创建Word文档
        doc = Document()

        # 添加标题
        title = doc.add_heading(excel_path.stem, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 提前收集所有图片信息（行号 -> 图片列表）
        images_by_row = {}
        if hasattr(ws, '_images'):
            for image in ws._images:
                if hasattr(image, 'anchor'):
                    row = image.anchor._from.row + 1
                    if row not in images_by_row:
                        images_by_row[row] = []
                    images_by_row[row].append(image)
            logger.info(f"   收集到图片行: {list(images_by_row.keys())}")

        # 逐行处理数据
        qa_count = 0
        for row_idx in range(2, ws.max_row + 1):
            # 读取行数据
            cells = [ws.cell(row_idx, col).value for col in range(1, 5)]

            # 过滤空行
            if all(c is None for c in cells):
                continue

            # 根据列数判断格式
            if isinstance(cells[0], (int, float)) and isinstance(cells[1], str):
                category = str(cells[3]) if cells[3] else ""
                seq_num = int(cells[0])
                question = str(cells[1])
                answer = str(cells[2]) if cells[2] else ""
            else:
                category = str(cells[0]) if cells[0] else ""
                seq_num = cells[1]
                question = str(cells[2]) if cells[2] else ""
                answer = str(cells[3]) if cells[3] else ""

            # 如果有图片，调用Vision API分析（无论答案是否存在）
            images_data = []  # 存储 (图片数据, 图片说明) 元组
            if row_idx in images_by_row:
                logger.info(f"   行{row_idx}: 检测到{len(images_by_row[row_idx])}张图片，使用Vision分析...")

                for idx, image in enumerate(images_by_row[row_idx]):
                    if hasattr(image, 'ref'):
                        try:
                            image.ref.seek(0)
                            image_data = image.ref.read()

                            logger.info(f"      分析图片{idx+1}/{len(images_by_row[row_idx])}...")

                            # 无论答案是否存在，都调用Vision API
                            # 如果无答案：返回详细内容作为答案 + 说明
                            # 如果有答案：返回空内容 + 说明（只生成图片说明）
                            full_content, description = analyze_image_with_question(
                                image_data,
                                question,
                                row_idx,
                                existing_answer=answer
                            )

                            # 有图片时，用图片解析内容补充答案
                            if full_content and full_content != "[分析失败]":
                                if not answer.strip():
                                    # 答案为空，直接用图片内容
                                    answer = full_content
                                    logger.info(f"      ✅ 生成答案")
                                elif full_content != answer:
                                    # 已有答案但图片有额外内容，追加到后面
                                    answer = answer.rstrip() + "\n\n" + full_content
                                    logger.info(f"      ✅ 追加图片解析内容")
                            else:
                                logger.info(f"      ✅ 完成")

                            images_data.append((image_data, description))
                            logger.info(f"      ✅ 完成")
                        except Exception as e:
                            logger.warning(f"      图片处理异常: {str(e)}")
                            images_data.append((image_data, "[处理异常]"))

            # 后续使用images_data代替images_by_row[row_idx]


            # 只有当问题为空时才跳过此行
            if not question.strip():
                if row_idx in images_by_row:
                    for image in images_by_row[row_idx]:
                        try:
                            if hasattr(image, 'ref'):
                                img_para = doc.add_paragraph()
                                image.ref.seek(0)
                                img_para.add_run().add_picture(image.ref, width=Inches(4.5))
                        except Exception as e:
                            pass
                continue

            # 添加问题
            if category.strip():
                question_text = f"【{category}】Q{seq_num}: {question}"
            else:
                question_text = f"Q{seq_num}: {question}"
            doc.add_heading(question_text, level=3)

            # 处理答案和图片
            # 先添加原始答案或Vision生成的答案
            if answer.strip():
                # 判断是否是markdown代码块格式的详细解析
                if '```markdown' in answer or '```' in answer:
                    # 整个markdown块作为一个段落添加
                    doc.add_paragraph(f"答: {answer}")
                else:
                    # 普通文本，逐行添加
                    answer_lines = answer.split('\n')
                    answer_lines = [line.strip() for line in answer_lines if line.strip()]
                    for line_idx, line in enumerate(answer_lines):
                        if line_idx == 0:
                            doc.add_paragraph(f"答: {line}")
                        else:
                            doc.add_paragraph(line)

            # 添加图片和图片说明
            if images_data:
                for img_idx, (image_data_bytes, image_title) in enumerate(images_data):
                    try:
                        # 添加图片
                        p = doc.add_paragraph()
                        p.add_run().add_picture(io.BytesIO(image_data_bytes), width=Inches(4.5))

                        # 添加图片标题说明
                        if image_title and image_title not in ["[处理异常]", "[分析失败]"]:
                            if len(images_data) > 1:
                                # 多张图片时添加编号
                                title_text = f"图片{img_idx+1}: {image_title}"
                            else:
                                title_text = f"图片说明: {image_title}"
                            doc.add_paragraph(title_text, style='List Bullet')
                    except Exception as e:
                        logger.warning(f"   行{row_idx}: 图片{img_idx+1}插入失败 - {str(e)}")

            # 问答之间的分隔符（两个空行）
            doc.add_paragraph()
            doc.add_paragraph()
            qa_count += 1

        # 保存Word文档
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        file_size = output_path.stat().st_size / 1024
        logger.info(f"✅ 成功: {output_path.name} ({qa_count} 条Q&A, {file_size:.1f}KB)")

        return True

    except Exception as e:
        logger.error(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """主函数 - 命令行参数版

    用法示例：
        # 转换单个文件（输出到同目录，自动替换.xlsx为.docx）
        python excel_to_word.py 知识库base_cw_0410/差旅智能体_大模型数据-知识库.xlsx

        # 转换单个文件并指定输出路径
        python excel_to_word.py input.xlsx output.docx

        # 批量转换目录下所有xlsx文件
        python excel_to_word.py 知识库base_cw_0410/

        # 转换多个指定文件
        python excel_to_word.py 知识库base_cw_0410/差旅智能体_大模型数据-知识库.xlsx 知识库base_cw_0410/差旅智能体_差旅标准问答.xlsx
    """
    import argparse

    parser = argparse.ArgumentParser(
        description="Excel to Word 转换工具 - 使用Vision API分析图片填充答案"
    )
    parser.add_argument(
        "inputs",
        nargs="+",
        help="单个xlsx文件路径 / 输出docx路径 / 要批量转换的目录"
    )

    args = parser.parse_args()

    # 收集要转换的文件列表
    files = []  # (input_path, output_path)
    for input_arg in args.inputs:
        path = Path(input_arg)
        if path.is_dir():
            # 目录下所有xlsx文件
            for xlsx in sorted(path.glob("*.xlsx")):
                output = xlsx.with_suffix(".docx")
                files.append((str(xlsx), str(output)))
        elif path.exists():
            if path.suffix == ".xlsx":
                # 单个文件，输出到同目录
                output = path.with_suffix(".docx")
                files.append((str(path), str(output)))
            else:
                logger.warning(f"⚠️  跳过非xlsx文件: {path}")
        else:
            logger.warning(f"⚠️  文件不存在: {path}")

    if not files:
        logger.error("❌ 没有找到要转换的文件")
        return

    logger.info("="*60)
    logger.info("Excel to Word 转换工具 (document类型Vision分析)")
    logger.info("="*60 + "\n")
    logger.info(f"待处理: {len(files)} 个文件\n")

    success_count = 0
    for input_file, output_file in files:
        if convert_excel_to_word(input_file, output_file):
            success_count += 1
        print()

    logger.info("="*60)
    logger.info(f"完成: {success_count}/{len(files)} 个文件转换成功")
    logger.info("="*60)


if __name__ == "__main__":
    main()
